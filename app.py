from flask import Flask, render_template, request, send_file, jsonify
import re, os, json, traceback
from io import BytesIO
from datetime import datetime

APP_VERSION = "v3.5-dist-letras"

# --- Dependencias opcionales para DOCX ---
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# --- Constantes de app ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, "static", "logo_hc.png")
HEADER_TEXT = "Este programa fue creado por Honduras Constructores S de R L"

app = Flask(__name__)

# --- Utilidades para textos en español ---
UNIDADES = ["cero","uno","dos","tres","cuatro","cinco","seis","siete","ocho","nueve"]
ESPECIALES_10_19 = ["diez","once","doce","trece","catorce","quince","dieciséis","diecisiete","dieciocho","diecinueve"]
VEINTI = ["veinte","veintiuno","veintidós","veintitrés","veinticuatro","veinticinco","veintiséis","veintisiete","veintiocho","veintinueve"]
DECENAS = [None,None,"veinte","treinta","cuarenta","cincuenta","sesenta","setenta","ochenta","noventa"]
CIENTOS = {100:"cien",200:"doscientos",300:"trescientos",400:"cuatrocientos",500:"quinientos",600:"seiscientos",700:"setecientos",800:"ochocientos",900:"novecientos"}

def numero_a_palabras(n:int)->str:
    """0–999"""
    if n<0 or n>999: raise ValueError("Solo se admite 0–999 para esta app")
    if n<10: return UNIDADES[n]
    if 10<=n<=19: return ESPECIALES_10_19[n-10]
    if 20<=n<=29: return VEINTI[n-20]
    if 30<=n<=99:
        d,u=divmod(n,10)
        return DECENAS[d] if u==0 else f"{DECENAS[d]} y {UNIDADES[u]}"
    c=(n//100)*100; r=n%100
    if n==100: return CIENTOS[100]
    pref=CIENTOS.get(c,"ciento")
    return pref if r==0 else f"{pref} {numero_a_palabras(r)}"

def entero_a_palabras_miles(n:int)->str:
    """0–999,999 usando numero_a_palabras para cada bloque de 0–999."""
    if n<1000:
        return numero_a_palabras(n)
    miles, resto = divmod(n, 1000)
    if miles == 1:
        pref = "mil"
    else:
        pref = f"{numero_a_palabras(miles)} mil"
    return pref if resto==0 else f"{pref} {numero_a_palabras(resto)}"

def forma_masculina(frase:str)->str:
    # uno -> un / veintiuno -> veintiún / "y uno" -> "y un"
    frase=re.sub(r"\bveintiuno\b","veintiún",frase)
    frase=re.sub(r" y uno\b"," y un",frase)
    frase=re.sub(r"\buno\b","un",frase)
    return frase

def plural_si_corresponde(v:int,sing:str)->str:
    return sing if v==1 else sing+"s"

ETIQUETA_RE=re.compile(r"^(?P<num>\d+)?(?P<letras>[A-Za-z]+)?$")

def etiqueta_a_texto(etq:str, convertir_numeros=True)->str:
    etq=(etq or "").strip()
    m=ETIQUETA_RE.match(etq)
    if not m: return etq
    num=m.group("num"); letras=(m.group("letras") or "").upper()
    partes=[]
    if num: partes.append(numero_a_palabras(int(num)) if convertir_numeros else num)
    if letras: partes.append(letras)
    return " ".join(partes) if partes else etq

CARD_WORD={"N":"Norte","S":"Sur","E":"Este","O":"Oeste","W":"Oeste"}

def parsear_rumbo_texto(raw:str):
    """Acepta: 
       - 'N, 25, 35, 20, O'
       - 'S 10°0\'30\'\' E'
       - 'N 10 5 0 O'
       - 'Norte ... Oeste' y también 'W' como 'O'.
    """
    if not raw: return None
    t=raw.strip()
    if not t: return None
    # Palabras→letras
    t = re.sub(r'\bNORTE\b','N',t,flags=re.I)
    t = re.sub(r'\bSUR\b','S',t,flags=re.I)
    t = re.sub(r'\bESTE\b','E',t,flags=re.I)
    t = re.sub(r'\bOESTE\b','O',t,flags=re.I)
    # Normalizaciones
    norm=t.replace("°"," ").replace("º"," ").replace("’","'").replace("´","'")
    norm=re.sub(r"[;|/]+"," ",norm)
    norm = re.sub(r'\bW\b','O',norm,flags=re.I)
    # Con comas
    partes=[p.strip() for p in norm.split(',') if p.strip()]
    if len(partes)>=5:
        c1=partes[0].upper()
        try:
            g=int(re.sub(r"\D","", partes[1] or "0"))
            m=int(re.sub(r"\D","", partes[2] or "0"))
            s=int(re.sub(r"\D","", partes[3] or "0"))
        except ValueError:
            return None
        c2=partes[4].upper()
        return (c1,g,m,s,c2)
    # Libre
    cards=re.findall(r"[NnSsEeOoWw]",norm)
    nums=re.findall(r"\d+",norm)
    if len(nums)>=3 and len(cards)>=2:
        c1=cards[0].upper(); c2=cards[-1].upper()
        if c1=='W': c1='O'
        if c2=='W': c2='O'
        g=int(nums[0]); m=int(nums[1]); s=int(nums[2])
        return (c1,g,m,s,c2)
    return None

def rumbo_texto(card1:str,g:int,m:int,s:int,card2:str)->str:
    g_w=forma_masculina(numero_a_palabras(g))
    m_w=forma_masculina(numero_a_palabras(m))
    s_w=forma_masculina(numero_a_palabras(s))
    return (f"{CARD_WORD.get(card1,card1)} {g_w} {plural_si_corresponde(g,'grado')}, "
            f"{m_w} {plural_si_corresponde(m,'minuto')}, "
            f"{s_w} {plural_si_corresponde(s,'segundo')} {CARD_WORD.get(card2,card2)}")

def rumbo_compacto_usuario(card1:str,g:int,m:int,s:int,card2:str)->str:
    """Devuelve exactamente: N 25° 35´20´´O (minutos y segundos sin espacio; último cardinal pegado)."""
    return f"{card1} {g}° {m}´{s}´´{card2}"

def normalizar_colindancia(txt:str)->str:
    """Si el usuario no escribe 'Colinda con', se antepone; si ya lo puso, no se duplica."""
    if not txt: return ""
    t = txt.strip()
    if not t: return ""
    if re.match(r"(?i)^\s*colinda\s+con\b", t):
        return t[0].upper() + t[1:]  # capitaliza la primera
    return "Colinda con " + t

def distancia_a_palabras(distancia_raw:str)->str|None:
    """
    Convierte '10.15' -> 'diez punto quince metros'
    Maneja ',' como separador decimal. Soporta enteros hasta 999,999.
    """
    if not distancia_raw: return None
    txt = distancia_raw.strip().replace(",", ".")
    # Solo números positivos con opcional decimal
    if not re.fullmatch(r"\d+(?:\.\d+)?", txt):
        return None
    int_part_str, dot, frac_part_str = txt.partition(".")
    try:
        int_val = int(int_part_str)
    except ValueError:
        return None
    # 0–999,999
    if int_val > 999_999:
        # Si es muy grande, no lo convertimos a palabras
        int_words = int_part_str  # fallback: dígitos
    else:
        int_words = entero_a_palabras_miles(int_val)

    if dot and frac_part_str:
        # quitar ceros a la izquierda en la parte decimal para pronunciar natural
        frac_trim = frac_part_str.lstrip("0")
        if frac_trim == "":
            # si era 10.00 -> solo 'diez metros'
            return f"{int_words} metros"
        try:
            frac_val = int(frac_trim)
        except ValueError:
            return f"{int_words} metros"
        if frac_val <= 999:
            frac_words = numero_a_palabras(frac_val)
        else:
            frac_words = frac_trim  # fallback a dígitos si excede 999
        return f"{int_words} punto {frac_words} metros"
    else:
        return f"{int_words} metros"

# ---------- Ayudante para construir tramos (lo usa index y descargar) ----------
def construir_tramos_desde_form(form):
    convertir = form.get('convertir','on')=='on'
    est_ini_list = form.getlist('est_ini[]')
    est_fin_list = form.getlist('est_fin[]')
    rumbo_txt_list = form.getlist('rumbo_texto[]')
    distancia_list = form.getlist('distancia[]')
    colind_list = form.getlist('colindancia[]')

    n = max(len(est_ini_list), len(est_fin_list), len(rumbo_txt_list), len(distancia_list), len(colind_list))
    tramos = []
    errores = []
    prev_fin_raw = None

    for i in range(n):
        est_ini = (est_ini_list[i] if i < len(est_ini_list) else '').strip()
        est_fin = (est_fin_list[i] if i < len(est_fin_list) else '').strip()
        rumbo_raw = (rumbo_txt_list[i] if i < len(rumbo_txt_list) else '').strip()
        distancia_raw = (distancia_list[i] if i < len(distancia_list) else '').strip()
        colind = (colind_list[i] if i < len(colind_list) else '').strip()

        # Auto-encadenar: si no hay inicio y existe un fin previo, usarlo
        if not est_ini and prev_fin_raw:
            est_ini = prev_fin_raw

        # Saltar filas completamente vacías
        if not (est_ini or est_fin or rumbo_raw or distancia_raw or colind):
            continue

        if not est_ini or not est_fin:
            errores.append(f"Fila {i+1}: ingresa estación inicio y fin (auto-encadené inicio='{est_ini or '∅'}').")
            prev_fin_raw = est_fin or prev_fin_raw
            continue

        parsed = parsear_rumbo_texto(rumbo_raw)
        if not parsed:
            errores.append(f"Fila {i+1}: no pude interpretar el rumbo \"{rumbo_raw}\".")
            prev_fin_raw = est_fin
            continue
        c1,g,m,s,c2 = parsed

        # Distancia (num) y en letras
        distancia = None
        dist_letras = None
        if distancia_raw:
            try:
                distancia = float(distancia_raw.replace(",", "."))
            except ValueError:
                errores.append(f"Fila {i+1}: distancia inválida.")
            dist_letras = distancia_a_palabras(distancia_raw)

        est_ini_txt = etiqueta_a_texto(est_ini, convertir_numeros=convertir)
        est_fin_txt = etiqueta_a_texto(est_fin, convertir_numeros=convertir)
        texto_rumbo = rumbo_texto(c1,g,m,s,c2)
        compacto_usuario = rumbo_compacto_usuario(c1,g,m,s,c2)
        colind_fmt = normalizar_colindancia(colind)

        redaccion = (f"De la estación {est_ini_txt} a la estación {est_fin_txt}, "
                     f"con rumbo {texto_rumbo} ({compacto_usuario}).")
        if distancia is not None:
            if dist_letras:
                redaccion += f" Distancia {distancia:.2f} m ({dist_letras})."
            else:
                redaccion += f" Distancia {distancia:.2f} m."
        if colind_fmt:
            redaccion += f" {colind_fmt}"

        tramos.append({
            "est_ini_txt": est_ini_txt,
            "est_fin_txt": est_fin_txt,
            "rumbo_texto": texto_rumbo,
            "rumbo_compacto": compacto_usuario,
            "distancia": distancia,
            "distancia_letras": dist_letras,
            "colindancia": colind_fmt,
            "redaccion": redaccion
        })

        prev_fin_raw = est_fin

    return tramos, errores

# ---------- Rutas ----------
@app.route('/_version')
def version():
    return jsonify({"version": APP_VERSION, "docx": DOCX_AVAILABLE})

@app.route('/', methods=['GET','POST'])
def index():
    errores=[]; resultado=None
    if request.method=='POST':
        tramos, errores = construir_tramos_desde_form(request.form)
        if tramos and not errores:
            redaccion_total = "\n".join(f"{idx+1}) {t['redaccion']}" for idx, t in enumerate(tramos))
            resultado = {"tramos": tramos, "redaccion_total": redaccion_total}
        elif not tramos and not errores:
            errores.append("Agrega al menos un tramo.")

    return render_template('formulario.html',
                           errores=errores,
                           resultado=resultado,
                           docx_ready=DOCX_AVAILABLE,
                           app_version=APP_VERSION)

@app.route('/descargar', methods=['POST'])
def descargar():
    if not DOCX_AVAILABLE:
        return "La librería python-docx no está instalada en el servidor.", 500

    # 1) Camino principal: payload_json desde el template
    tramos = None
    payload_json = request.form.get('payload_json', '')
    if payload_json:
        try:
            data = json.loads(payload_json)
            tramos = data.get("tramos", [])
        except Exception:
            tramos = None

    # 2) Fallback: reconstruir desde campos del formulario (por si tu template no envía payload_json)
    if tramos is None or not isinstance(tramos, list) or not tramos:
        tramos, errores = construir_tramos_desde_form(request.form)
        if errores and not tramos:
            return "No pude leer datos para el Word. Completa al menos un tramo.", 400

    try:
        doc = Document()
        styles=doc.styles['Normal']; styles.font.name='Calibri'; styles.font.size=Pt(11)

        # Encabezado con logo + texto
        section=doc.sections[0]; header=section.header
        table=header.add_table(rows=1, cols=2); table.autofit=True
        try:
            if os.path.exists(LOGO_PATH):
                left_p=table.cell(0,0).paragraphs[0]
                left_p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                left_p.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
        except Exception:
            pass
        right_p=table.cell(0,1).paragraphs[0]; right_p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        rt=right_p.add_run(HEADER_TEXT); rt.bold=True; rt.font.size=Pt(10)

        doc.add_heading('Redacción topográfica', level=1)
        doc.add_paragraph(datetime.now().strftime("Generado el %Y-%m-%d %H:%M:%S"))

        # Resumen
        doc.add_paragraph().add_run("Resumen de tramos:").bold = True
        for i, t in enumerate(tramos, start=1):
            doc.add_paragraph(f"{i}) {t['redaccion']}")

        # Detalle
        doc.add_paragraph().add_run("Detalle:").bold = True
        tbl = doc.add_table(rows=1, cols=7)
        hdr = tbl.rows[0].cells
        hdr[0].text = "Est. Inicio"
        hdr[1].text = "Est. Fin"
        hdr[2].text = "Rumbo (texto)"
        hdr[3].text = "Rumbo compacto"
        hdr[4].text = "Distancia (m)"
        hdr[5].text = "Distancia (letras)"
        hdr[6].text = "Colindancia"

        for t in tramos:
            row = tbl.add_row().cells
            row[0].text = t["est_ini_txt"]
            row[1].text = t["est_fin_txt"]
            row[2].text = t["rumbo_texto"]
            row[3].text = t["rumbo_compacto"]
            row[4].text = "" if t["distancia"] is None else f"{t['distancia']:.2f}"
            row[5].text = t.get("distancia_letras") or ""
            row[6].text = t["colindancia"] or ""

        bio = BytesIO(); doc.save(bio); bio.seek(0)
        return send_file(bio, as_attachment=True, download_name="Redaccion_Multitramos.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception:
        return "Error generando DOCX:\n" + traceback.format_exc(), 500

if __name__=='__main__':
    app.run(debug=True)
